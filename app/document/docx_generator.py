import io
from typing import Dict, Any, List
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from app.core.logging import logger
from app.core.exceptions import DocxGenerationException

# Hex color definitions
COLOR_PRIMARY_HEX = "1A365D"       # Deep Navy
COLOR_SECONDARY_HEX = "2B6CB0"     # Accent Blue
COLOR_TEXT_HEX = "2D3748"          # Charcoal slate
COLOR_BG_LIGHT_HEX = "F7FAFC"      # Light gray background for rows
COLOR_WHITE_HEX = "FFFFFF"

class DocxGenerator:
    """Generates professionally-styled Microsoft Word documents (.docx) from structured JSON."""
    
    def generate(self, data: Dict[str, Any]) -> bytes:
        """
        Translates structured document JSON into a styled DOCX byte stream.
        
        :param data: Dictionary conforming to schemas.content.StructuredDocument
        :return: raw bytes of the generated .docx file
        """
        logger.info("DocxGenerator: Initializing document generation.")
        try:
            doc = Document()
            
            # 1. Page Margins Setup
            for section in doc.sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
                
            # 2. Typography Styles Setup
            self._configure_styles(doc)
            
            # 3. Title Section
            title_text = data.get("title", "Enterprise Strategy Document")
            p_title = doc.add_paragraph(title_text, style="DocTitle")
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 4. Metadata Box / Subtitle
            meta = data.get("metadata", {})
            author = meta.get("author", "Autonomous Agent")
            date_str = datetime.now().strftime("%Y-%m-%d")
            
            p_meta = doc.add_paragraph()
            p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_meta = p_meta.add_run(f"Prepared by: {author}  |  Date: {date_str}")  
            r_meta.font.italic = True
            r_meta.font.size = Pt(10)
            r_meta.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
            
            # Add spacing below title block
            doc.add_paragraph().paragraph_format.space_after = Pt(24)
            
            # 5. Header & Footer Setup
            self._configure_header_footer(doc, title_text, date_str)
            
            # 6. Iterate and Add Sections
            sections = data.get("sections", [])
            for idx, sec in enumerate(sections):
                heading = sec.get("heading")
                subheading = sec.get("subheading")
                paragraphs = sec.get("paragraphs", [])
                bullets = sec.get("bullets", [])
                tables = sec.get("tables", [])
                references = sec.get("references", [])
                
                # Heading 1
                if heading:
                    p_h1 = doc.add_paragraph(heading, style="DocHeading1")
                    # Keep heading with next paragraph to prevent orphaned headings
                    p_h1.paragraph_format.keep_with_next = True
                    
                # Heading 2 / Subheading
                if subheading:
                    p_h2 = doc.add_paragraph(subheading, style="DocHeading2")
                    p_h2.paragraph_format.keep_with_next = True
                    
                # Paragraphs
                for para in paragraphs:
                    doc.add_paragraph(para, style="DocBody")
                    
                # Bulleted List
                for bullet in bullets:
                    doc.add_paragraph(bullet, style="List Bullet")
                    
                # Tables
                for table_data in tables:
                    self._add_styled_table(doc, table_data)
                    
                # References / Citations
                if references:
                    p_ref_title = doc.add_paragraph("References:", style="DocHeading2")
                    p_ref_title.paragraph_format.keep_with_next = True
                    for ref in references:
                        p_ref = doc.add_paragraph(ref, style="DocBody")
                        p_ref.paragraph_format.left_indent = Inches(0.25)
                        p_ref.runs[0].font.size = Pt(9.5)
                        p_ref.runs[0].font.italic = True
                        
                # Add spacing between sections
                if idx < len(sections) - 1:
                    p_spacer = doc.add_paragraph()
                    p_spacer.paragraph_format.space_after = Pt(12)
                    
            # 7. Write to stream
            stream = io.BytesIO()
            doc.save(stream)
            logger.info("DocxGenerator: Word document compilation complete.")
            return stream.getvalue()
            
        except Exception as e:
            logger.error(f"DocxGenerator: Generation failure: {e}")
            raise DocxGenerationException(f"Failed to generate Word document binary: {str(e)}")

    def _configure_styles(self, doc: Document) -> None:
        """Configures typography styles for title, headings, and body."""
        styles = doc.styles
        
        # Primary Font Settings
        FONT_FAMILY = "Calibri"
        
        # DocTitle Style
        style_title = styles.add_style("DocTitle", 1)
        font_title = style_title.font
        font_title.name = FONT_FAMILY
        font_title.size = Pt(24)
        font_title.bold = True
        font_title.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        style_title.paragraph_format.space_after = Pt(6)
        
        # DocHeading1 Style
        style_h1 = styles.add_style("DocHeading1", 1)
        font_h1 = style_h1.font
        font_h1.name = FONT_FAMILY
        font_h1.size = Pt(18)
        font_h1.bold = True
        font_h1.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        style_h1.paragraph_format.space_before = Pt(16)
        style_h1.paragraph_format.space_after = Pt(6)
        
        # DocHeading2 Style
        style_h2 = styles.add_style("DocHeading2", 1)
        font_h2 = style_h2.font
        font_h2.name = FONT_FAMILY
        font_h2.size = Pt(13)
        font_h2.bold = True
        font_h2.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        style_h2.paragraph_format.space_before = Pt(12)
        style_h2.paragraph_format.space_after = Pt(4)
        
        # DocBody Style (Normal override)
        style_body = styles["Normal"]
        font_body = style_body.font
        font_body.name = FONT_FAMILY
        font_body.size = Pt(11)
        font_body.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        style_body.paragraph_format.line_spacing = 1.15
        style_body.paragraph_format.space_after = Pt(6)
        
        # Explicit DocBody style alias for safety
        style_body_alias = styles.add_style("DocBody", 1)
        font_body_a = style_body_alias.font
        font_body_a.name = FONT_FAMILY
        font_body_a.size = Pt(11)
        font_body_a.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        style_body_alias.paragraph_format.line_spacing = 1.15
        style_body_alias.paragraph_format.space_after = Pt(6)

    def _configure_header_footer(self, doc: Document, title: str, date_str: str) -> None:
        """Styles the document running header and running footer with dynamic page numbers."""
        for section in doc.sections:
            # 1. Header (Right aligned, italic title)
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = ""
            header_run = header_para.add_run(f"Enterprise Knowledge Agent  |  {title}")
            header_run.font.name = "Calibri"
            header_run.font.size = Pt(8.5)
            header_run.font.italic = True
            header_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 2. Footer (Left: Date, Right: Page X of Y)
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = ""
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Left block - Date (Word handles alignment spacing via tabs or spaces in basic footers)
            date_run = footer_para.add_run(f"{date_str}       |       Confidential       |       Page ")
            date_run.font.name = "Calibri"
            date_run.font.size = Pt(8.5)
            date_run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
            
            # Dynamic PAGE Number field injection
            self._inject_page_number_field(footer_para.add_run())

    def _inject_page_number_field(self, run) -> None:
        """Injects dynamic MS Word page number XML field codes into a run."""
        fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
        fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
        fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    def _add_styled_table(self, doc: Document, table_data: Dict[str, Any]) -> None:
        """Renders and styles a table structure with colored headers and alternating rows."""
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        
        if not headers:
            return
            
        # Create Table in Document
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        # Format Headers
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = str(text)
            self._set_cell_background(hdr_cells[idx], COLOR_PRIMARY_HEX)
            
            # Set white bold font for headers
            for p in hdr_cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    
        # Format Rows
        for r_idx, row_items in enumerate(rows):
            row_cells = table.add_row().cells
            bg_color = COLOR_BG_LIGHT_HEX if r_idx % 2 == 1 else COLOR_WHITE_HEX
            
            for c_idx, cell_text in enumerate(row_items):
                if c_idx < len(row_cells):
                    row_cells[c_idx].text = str(cell_text)
                    if bg_color != COLOR_WHITE_HEX:
                        self._set_cell_background(row_cells[c_idx], bg_color)
                        
                    # Set body text styles
                    for p in row_cells[c_idx].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in p.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(9.5)
                            run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
                            
        # Append spacing below table
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    def _set_cell_background(self, cell, fill_hex: str) -> None:
        """Manipulates underlying OpenXML of cell to set background shading."""
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

# Global generator instance
docx_generator = DocxGenerator()
