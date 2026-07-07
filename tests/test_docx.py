import pytest
from app.document.docx_generator import docx_generator
from docx import Document
import io

def test_docx_generation_formatting():
    # Input data matching StructuredDocument format
    doc_data = {
        "title": "Enterprise Technical Specification",
        "metadata": {
            "author": "Chief Architect",
            "version": "1.4.2",
            "date": "July 07, 2026"
        },
        "sections": [
            {
                "heading": "1. Architectural Overview",
                "subheading": "1.1 Client-Server Communication",
                "paragraphs": [
                    "This paragraph outlines the main communication boundaries of the decoupled application.",
                    "The client executes async requests to FastAPI endpoints."
                ],
                "bullets": [
                    "HTTP/2 protocol constraints",
                    "Keep-alive TCP pooling"
                ],
                "tables": [
                    {
                        "headers": ["Service", "Port", "Access"],
                        "rows": [
                            ["API Gateway", "8000", "Public"],
                            ["Vector Database", "8500", "Private"]
                        ]
                    }
                ],
                "references": ["IEEE 802.3 Ethernet Standards"]
            }
        ]
    }
    
    # Generate binary
    docx_bytes = docx_generator.generate(doc_data)
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 0
    
    # Reload document using python-docx to verify content is present and structurally correct
    doc = Document(io.BytesIO(docx_bytes))
    
    # Verify Title
    assert doc.paragraphs[0].text == "Enterprise Technical Specification"
    assert doc.paragraphs[0].style.name == "DocTitle"
    
    # Verify metadata paragraph
    assert "Prepared by: Chief Architect" in doc.paragraphs[1].text
    
    # Verify Headings
    h1_texts = [p.text for p in doc.paragraphs if p.style.name == "DocHeading1"]
    assert "1. Architectural Overview" in h1_texts
    
    h2_texts = [p.text for p in doc.paragraphs if p.style.name == "DocHeading2"]
    assert "1.1 Client-Server Communication" in h2_texts
    assert "References:" in h2_texts
    
    # Verify Bullets
    bullet_texts = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert "HTTP/2 protocol constraints" in bullet_texts
    
    # Verify Table
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # Check headers
    assert table.rows[0].cells[0].text == "Service"
    assert table.rows[0].cells[1].text == "Port"
    assert table.rows[0].cells[2].text == "Access"
    # Check rows
    assert table.rows[1].cells[0].text == "API Gateway"
    assert table.rows[2].cells[0].text == "Vector Database"
